from docx import Document
from docx.shared import Pt

def create_homework_doc():
    doc = Document()

    # Titolo
    doc.add_heading('Calcolo Numerico 2025-26', 0)
    doc.add_heading('Homework 1: Zeri di funzione', 1)

    # Consegna
    doc.add_heading('Consegna', 2)
    p = doc.add_paragraph('Zeri di funzione: utilizzare il metodo di bisezione, il metodo delle iterazioni di punto fisso e il metodo di Newton per il calcolo dello zero delle seguenti funzioni.')
    doc.add_paragraph('Disegnare il grafico della funzione per localizzare lo zero e scegliere sia l\'intervallo in cui applicare la bisezione che il punto iniziale per gli altri due algoritmi.')

    # Funzioni (Lista)
    functions = [
        "1. f(x) = ln(x+1) - x    (g(x) = ln(x+1))",
        "2. f(x) = x^2 - cos(x)   (g(x) = sqrt(cos(x)))",
        "3. f(x) = sin(x) - x^2   (g(x) = 2*sin(x))",
        "4. f(x) = e^x - 3x       (g(x) = (1/3)e^x)"
    ]
    for f in functions:
        doc.add_paragraph(f)

    # Codici
    doc.add_page_break()
    doc.add_heading('Codici degli algoritmi', 2)
    
    codes = {
        "Metodo di Bisezione": """def metodo_bisezione(f, a, b, max_iter, epsilon):
    cond = True
    i = 0
    if f(a)*f(b) >= 0.0:
        return None, 0
    while cond and i < max_iter:
        c = (a+b)/2
        if abs(f(c)) < epsilon:
            cond = False
        if cond and f(a)*f(c) < 0:
            b = c
        elif cond:
            a = c
        i = i + 1
    return c, i""",
        
        "Metodo Punto Fisso": """def metodo_punto_fisso(g, x0, max_iter, epsilon1, epsilon2):
    i = 0
    cond = True
    xk = x0
    while cond and i < max_iter:
        x = g(xk)
        if abs(x - xk) < epsilon1: # Corretto da OCR
            cond = False
        xk = x
        i = i + 1
    return xk, i""",

        "Metodo di Newton": """def metodo_newton(f, df, x0, max_iter, epsilon1, epsilon2):
    i = 0
    cond = True
    xk = x0
    if abs(df(xk)) < 1e-10:
        return None, 0
    while cond and i < max_iter:
        xk1 = xk - f(xk) / df(xk)
        if abs(xk1 - xk) < epsilon1 or abs(f(xk1)) < epsilon2:
            cond = False
        xk = xk1
        i = i + 1
    return xk, i"""
    }

    for title, code in codes.items():
        doc.add_heading(title, 3)
        p = doc.add_paragraph(code)
        # Imposta monospace per il codice
        for run in p.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)

    # Tabelle Risultati
    doc.add_page_break()
    doc.add_heading('Risultati Esecuzioni (Esempio Funzione 2)', 2)
    doc.add_paragraph("f(x) = x^2 - cos(x)")

    # Creazione Tabelle (Dati dal PDF)
    data_sets = [
        [("Metodo", "Punto Iniziale", "Intervallo", "Soluzione", "Iterazioni"),
         ("Bisezione", "0", "[0, 1]", "0.824131...", "20"),
         ("Punto Fisso", "0", "[0, 1]", "0.824132...", "18"),
         ("Newton", "0", "[0, 1]", "N/A", "")],
         
        [("Metodo", "Punto Iniziale", "Intervallo", "Soluzione", "Iterazioni"),
         ("Bisezione", "0.5", "[0, 1]", "0.824131...", "20"),
         ("Punto Fisso", "0.5", "0.1", "0.824132...", "17"),
         ("Newton", "0.5", "[0, 1]", "0.824132...", "4")],
         
        [("Metodo", "Punto Iniziale", "Intervallo", "Soluzione", "Iterazioni"),
         ("Bisezione", "0.8", "[0, 1]", "0.824131...", "20"),
         ("Punto Fisso", "0.8", "[0, 1]", "0.824132...", "14"),
         ("Newton", "0.8", "[0, 1]", "0.824132...", "2")]
    ]

    for data in data_sets:
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(data[0]):
            hdr_cells[i].text = header
        
        for row_data in data[1:]:
            row_cells = table.add_row().cells
            for i, item in enumerate(row_data):
                row_cells[i].text = item
        doc.add_paragraph("\n")

    doc.save('Homework1_Zeri_Funzione.docx')
    print("File Word creato con successo!")

create_homework_doc()